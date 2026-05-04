import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_report():
    doc = Document()
    
    # Header: DONE BY
    p_done = doc.add_paragraph()
    p_done.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p_done.add_run("DONE BY :\nAMMOGH ASHOK    1RUB25CSE0001\nLAHARI PANDITH M   1RUB25CSE011")
    run.font.size = Pt(12)
    run.bold = True
    
    doc.add_paragraph() # spacing
    
    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p_title.add_run("Project Report: SNMP GetRequest Implementation")
    run.font.size = Pt(16)
    run.bold = True
    
    doc.add_paragraph() # spacing

    with open("report.md", "r", encoding="utf-8") as f:
        md_text = f.read()
        
    # Simple markdown parser for docx
    lines = md_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
            
        if line.startswith('# '):
            pass # Skip title as we already added it
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=2)
        elif line.startswith('* **') or line.startswith('- **'):
            # bullet with bold
            p = doc.add_paragraph(style='List Bullet')
            clean_line = line[4:] # remove '* **'
            parts = clean_line.split('**:')
            if len(parts) > 1:
                run = p.add_run(parts[0] + ":")
                run.bold = True
                p.add_run(parts[1])
            else:
                p.add_run(clean_line.replace('**', ''))
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('> '):
            p = doc.add_paragraph(line[2:])
            p.runs[0].italic = True
        elif line.startswith('**'):
            p = doc.add_paragraph()
            clean_line = line.replace('**', '')
            run = p.add_run(clean_line)
            run.bold = True
        else:
            doc.add_paragraph(line)

    doc.save("Project_Report_SNMP.docx")
    print("Project_Report_SNMP.docx created successfully!")

if __name__ == "__main__":
    create_report()
